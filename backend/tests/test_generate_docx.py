# backend/tests/test_generate_docx.py
#
# Tests for the document generator in generate_docx.py.
#
# These tests build a real .docx file from a minimal TailoredResumeOutput and
# check that the output is valid and contains the content we gave it.
#
# We don't check visual formatting (fonts, spacing, colors) — that would
# require inspecting raw XML and would be very brittle. Instead we focus on:
#   - The output file is actually created on disk.
#   - python-docx can open it without raising an error (i.e. it's valid OOXML).
#   - The summary text we passed in appears somewhere in the document.
#   - The function returns the output path (so FastAPI can hand it to FileResponse).

import pytest
import tempfile
from pathlib import Path
from docx import Document

from schemas.resume import TailoredResumeOutput, ExperienceEntry, SkillCategory
from generate_docx import generate_resume_docx


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def minimal_resume():
    """
    A minimal but valid TailoredResumeOutput.

    The document generator also writes hardcoded header and footer sections
    (name, contact, education, etc.) that don't come from this object, so
    we only need to supply the three variable fields: summary, experience,
    and skills.
    """
    return TailoredResumeOutput(
        summary="A Python engineer with strong API development skills.",
        experience=[
            ExperienceEntry(
                company="Acme Corp",
                title="Software Engineer",
                start_date="Jan 2023",
                end_date="Present",
                bullets=["Built REST APIs", "Improved CI/CD pipeline"],
            )
        ],
        skills=[
            SkillCategory(category="Languages", entries=["Python", "JavaScript"]),
        ],
    )


@pytest.fixture
def output_path():
    """
    A temporary file path for the generated DOCX.

    tempfile.NamedTemporaryFile creates an empty file so we have a path to
    pass to the generator. We close it immediately so python-docx can write
    to it on Windows (which locks open file handles). The `yield` gives the
    path to the test, then the cleanup block deletes the file afterwards so
    we don't leave temp files lying around.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = Path(f.name)
    yield path
    if path.exists():
        path.unlink()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestGenerateResumeDocx:
    def test_output_file_is_created(self, minimal_resume, output_path):
        """The function should write a file to the path it was given."""
        generate_resume_docx(minimal_resume, output_path)
        assert output_path.exists()

    def test_output_is_valid_docx(self, minimal_resume, output_path):
        """
        python-docx should be able to open the file without raising an error.

        A .docx file is a ZIP archive containing XML files. If the archive is
        malformed or the XML is invalid, Document() will raise an exception.
        We also check there's at least one paragraph — an empty document would
        be a sign something went badly wrong.
        """
        generate_resume_docx(minimal_resume, output_path)
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0

    def test_summary_text_appears_in_document(self, minimal_resume, output_path):
        """
        The summary string we provided should appear somewhere in the document.

        We join all paragraph text into one string and check for a distinctive
        phrase from the summary. This confirms the generator is actually writing
        the resume data we passed in, not just the hardcoded header.
        """
        generate_resume_docx(minimal_resume, output_path)
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Python engineer" in all_text

    def test_experience_company_appears_in_document(self, minimal_resume, output_path):
        """The company name from the experience entry should appear in the document."""
        generate_resume_docx(minimal_resume, output_path)
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text

    def test_returns_output_path(self, minimal_resume, output_path):
        """
        The function should return the path it was given.

        FastAPI's /generate endpoint passes the return value straight to
        FileResponse, so the function must return the path — not None.
        """
        result = generate_resume_docx(minimal_resume, output_path)
        assert result == output_path
