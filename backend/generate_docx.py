# backend/generate_docx.py
#
# Builds a formatted DOCX resume from a TailoredResumeOutput object.
#
# The core logic lives in generate_resume_docx(), which the FastAPI endpoint
# calls directly. The if __name__ == "__main__" block at the bottom preserves
# the original standalone behaviour so you can still run:
#
#   python generate_docx.py
#
# to generate a sample document from base_resume.json for visual inspection.

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.resume import TailoredResumeOutput

# ---------------------------------------------------------------------------
# Style constants — tweak these to change the look of the whole document
# ---------------------------------------------------------------------------
FONT_NAME  = "Calibri"
BODY_SIZE  = Pt(10)
SMALL_SIZE = Pt(9.5)

BLACK      = RGBColor(0x1A, 0x1A, 0x1A)   # near-black for body text
DARK_GRAY  = RGBColor(0x55, 0x55, 0x55)   # company / date lines

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _fmt_run(run, size=None, bold=False, italic=False, color=None):
    """
    Apply consistent font settings to a run (an inline span of text).
    All arguments are optional — only the ones passed will be set.
    """
    run.font.name  = FONT_NAME
    run.font.size  = size or BODY_SIZE
    run.bold       = bold
    run.italic     = italic
    run.font.color.rgb = color or BLACK


def _fmt_para(para, space_before=0, space_after=3):
    """
    Set paragraph spacing (in points) and disable the default space
    Word adds between lines inside a paragraph.
    """
    fmt = para.paragraph_format
    fmt.space_before    = Pt(space_before)
    fmt.space_after     = Pt(space_after)
    fmt.line_spacing    = Pt(13)   # ~single-spaced at 10pt


def _add_bottom_border(para):
    """
    Draw a thin horizontal rule under a paragraph.
    python-docx doesn't expose paragraph borders natively, so we write
    the required Office XML by hand:
      w:pBdr  = paragraph border container
      w:bottom = the bottom edge of that border
    """
    pPr    = para._p.get_or_add_pPr()       # get/create paragraph properties
    pBdr   = OxmlElement("w:pBdr")          # border container
    bottom = OxmlElement("w:bottom")        # bottom edge
    bottom.set(qn("w:val"),   "single")     # solid single line
    bottom.set(qn("w:sz"),    "4")          # 4 half-points = 0.5 pt thick
    bottom.set(qn("w:space"), "2")          # small gap between text and line
    bottom.set(qn("w:color"), "AAAAAA")     # light gray
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(doc, text):
    """Bold all-caps label with a light gray rule underneath."""
    p = doc.add_paragraph()
    _fmt_para(p, space_before=10, space_after=4)
    _add_bottom_border(p)
    _fmt_run(p.add_run(text), bold=True)
    return p


def add_bullet(doc, text):
    """
    Indented bullet point.
    Uses a hanging indent so the bullet character sits to the left and
    the text wraps neatly beneath itself (not beneath the bullet).
    """
    p = doc.add_paragraph()
    _fmt_para(p, space_before=0, space_after=2)
    fmt = p.paragraph_format
    fmt.left_indent        = Inches(0.25)
    fmt.first_line_indent  = Inches(-0.18)  # pulls the bullet left of the text
    _fmt_run(p.add_run(f"\u2013  {text}"), size=SMALL_SIZE)  # en-dash bullet
    return p



# ---------------------------------------------------------------------------
# Main function
# ---------------------------------------------------------------------------

def generate_resume_docx(resume: TailoredResumeOutput, output_path: Path) -> Path:
    """
    Build and save a DOCX resume from a TailoredResumeOutput object.

    Args:
        resume:      The assembled resume data (AI summary + passthrough content).
        output_path: Where to save the .docx file.

    Returns:
        output_path, so callers can immediately pass it to FastAPI's FileResponse.
    """
    # ---------------------------------------------------------------------------
    # Build the document
    # ---------------------------------------------------------------------------
    doc = Document()

    # Set page margins — 0.75" all round gives more usable space than Word's 1" default.
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Every new Document() starts with one empty paragraph — remove it so we
    # don't get a blank line at the top of the output.
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ---------------------------------------------------------------------------
    # HARDCODED HEADER — name, title, and contact details never change
    # ---------------------------------------------------------------------------

    # Name — large and bold
    p = doc.add_paragraph()
    _fmt_para(p, space_before=0, space_after=2)
    _fmt_run(p.add_run("Nicholas Riegel"), size=Pt(22), bold=True)

    # Title line
    p = doc.add_paragraph()
    _fmt_para(p, space_before=0, space_after=2)
    _fmt_run(p.add_run("AI Software Engineer  |  Full-Stack Developer"), size=Pt(11))

    # Citizenship / location
    p = doc.add_paragraph()
    _fmt_para(p, space_before=0, space_after=2)
    _fmt_run(p.add_run("Citizenship: Switzerland  |  Location: Bern, Switzerland"), size=SMALL_SIZE, color=DARK_GRAY)

    # Contact line
    p = doc.add_paragraph()
    _fmt_para(p, space_before=0, space_after=10)
    _fmt_run(p.add_run("nicholaspriegel@gmail.com  |  linkedin.com/in/nicholas-riegel  |  github.com/Nicholas-Riegel"), size=SMALL_SIZE, color=DARK_GRAY)

    # -- Summary ----------------------------------------------------------------
    if resume.summary:
        p = doc.add_paragraph()
        _fmt_para(p, space_before=0, space_after=8)
        _fmt_run(p.add_run(resume.summary))

    # -- Technical Skills -------------------------------------------------------
    add_section_header(doc, "TECHNICAL SKILLS")

    for cat in resume.skills:
        # Each category occupies one paragraph: bold label, then normal-weight entries.
        p = doc.add_paragraph()
        _fmt_para(p, space_before=0, space_after=2)
        _fmt_run(p.add_run(cat.category + ":  "), bold=True)
        _fmt_run(p.add_run(", ".join(cat.entries)))

    # -- Work Experience --------------------------------------------------------
    add_section_header(doc, "WORK EXPERIENCE")

    for job in resume.experience:
        # Job title — bold
        p = doc.add_paragraph()
        _fmt_para(p, space_before=7, space_after=1)
        _fmt_run(p.add_run(job.title), bold=True)

        # Company, date range, and location on one line, in muted gray
        start = job.start_date or ""
        end   = job.end_date   or ""
        dates = f"{start}\u2013{end}" if (start or end) else ""
        parts = [job.company, dates, job.location or ""]
        line  = "   |   ".join(part for part in parts if part)  # skip empty parts
        p2 = doc.add_paragraph()
        _fmt_para(p2, space_before=0, space_after=3)
        _fmt_run(p2.add_run(line), size=SMALL_SIZE, color=DARK_GRAY)

        # Bullet points
        for bullet in job.bullets:
            add_bullet(doc, bullet)

    # ---------------------------------------------------------------------------
    # HARDCODED FOOTER SECTIONS — these never change
    # ---------------------------------------------------------------------------

    # -- Professional Development -----------------------------------------------
    add_section_header(doc, "PROFESSIONAL DEVELOPMENT")

    prof_dev = [
        ("Data Structures & Algorithms  |  Udemy  |  in progress",
         "Completing a comprehensive course covering Big O notation, recursion, sorting and searching algorithms, linked lists, trees, graphs, and dynamic programming."),
        ("Advanced German  |  Französische Kirche  |  Bern",
         "Studying and learning advanced German."),
    ]
    for title, description in prof_dev:
        p = doc.add_paragraph()
        _fmt_para(p, space_before=5, space_after=1)
        _fmt_run(p.add_run(title), bold=True)
        p2 = doc.add_paragraph()
        _fmt_para(p2, space_before=0, space_after=4)
        _fmt_run(p2.add_run(description))

    # -- Education --------------------------------------------------------------
    add_section_header(doc, "EDUCATION")

    p = doc.add_paragraph()
    _fmt_para(p, space_before=5, space_after=1)
    _fmt_run(p.add_run("University of Toronto"), bold=True)
    p2 = doc.add_paragraph()
    _fmt_para(p2, space_before=0, space_after=4)
    _fmt_run(p2.add_run("PhD — Philosophy"))

    # -- Natural Languages ------------------------------------------------------
    add_section_header(doc, "NATURAL LANGUAGES")

    languages = [
        "English — fluent",
        "German — intermediate (B1/B2)",
        "French — advanced intermediate (B2)",
    ]
    for lang in languages:
        p = doc.add_paragraph()
        _fmt_para(p, space_before=0, space_after=2)
        _fmt_run(p.add_run(lang))

    # ---------------------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------------------
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Standalone mode — lets you run `python generate_docx.py` to visually
# check the output without needing the full FastAPI server running.
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import json
    from schemas.resume import BaseResume

    data_path = Path(__file__).parent / "data" / "base_resume.json"
    with open(data_path) as f:
        raw = json.load(f)

    # Load as BaseResume first, then convert to TailoredResumeOutput.
    # Both types now share the same ExperienceEntry and SkillCategory types,
    # so no .model_dump() conversion is needed — the objects pass straight through.
    base   = BaseResume(**raw)
    resume = TailoredResumeOutput(
        summary    = base.summary,
        skills     = base.skills,
        experience = base.experience,
    )

    out = Path(__file__).parent / "output" / "resume_sample.docx"
    generate_resume_docx(resume, out)
    print(f"DOCX saved to {out}")